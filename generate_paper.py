# -*- coding: utf-8 -*-
"""生成数模竞赛论文 Word 文档"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import docx.oxml

ROOT = r'D:\Math Modeling\2026 C题'
FIG_DIR = os.path.join(ROOT, 'output', 'figures')
OUTPUT = os.path.join(ROOT, 'output')

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, align=None, size=None, font_name=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    else:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_figure(caption, filename):
    path = os.path.join(FIG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        add_para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, font_name='宋体')

# ========== 标题 ==========
add_para('校园社交平台"好友关系与社交价值"分析', bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=22, font_name='黑体')
add_para('2026年西安电子科技大学数学建模校内赛 C题', bold=False,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()

# ========== 摘要 ==========
add_heading_styled('摘要', level=1)
add_para(
    '本文针对校园社交平台"校园圈"的用户关系与信息传播问题，基于好友关系表、用户属性表和行为数据表，'
    '综合运用复杂网络分析、多维度相似度计算和独立级联传播模型，完成了社群发现、好友推荐、关键用户识别和精准推送策略四项任务。'
    '\n\n社群发现方面，采用 Louvain 算法从 200 用户、1198 条好友关系中识别出 7 个社群（模块度 Q=0.2586），'
    '提取了密度最大的 5 个社群进行特征画像，发现社群间无成员重叠但功能定位各有侧重。'
    '\n\n好友推荐方面，构建了包含共同好友 Jaccard、同年级、同专业类别、共同社团、同活跃时段和行为余弦相似度的六维加权模型，'
    '为目标用户 S11 推荐了 S149（得分 0.567）、S146（0.536）和 S131（0.535）三位最合适的好友。'
    '\n\n信息传播方面，建立了基于转发概率公式的独立级联模型，通过蒙特卡洛模拟（30 次取平均）筛选出科技类话题关键用户 S6，'
    '其 48 小时传播范围达 186 人，传播过程呈 S 型增长曲线。精准推送策略达到 191 人传播范围，验证了策略的有效性。'
    '\n\n关键词：社交网络分析；Louvain 社区发现；好友推荐；独立级联模型；信息传播'
)

# ========== 目录占位 ==========
add_heading_styled('目录', level=1)
add_para('（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）')

# ========== 1. 问题重述 ==========
add_heading_styled('1. 问题重述', level=1)
add_para(
    '某高校自研的"校园圈"社交平台自 2021 年上线以来，沉淀了 8000 余名有效用户。然而平台面临"沉默的大多数"困境：'
    '日活在开学两周后下滑，头部用户贡献大部分内容，信息传播存在"院系墙"和"社团墙"现象，好友推荐沿用简单的"共同好友计数"逻辑。'
    '本题要求基于附件数据中的好友关系表、用户属性表和行为数据表，解决以下四个问题：')
add_para(
    '问题一：识别社交网络中的社群结构，找出内部连接密度最大的 5 个社群，分析成员分布和功能定位的重叠性。\n'
    '问题二：分析好友之间存在的共有特性，为目标用户 S11 推荐 3 位最合适的好友，并分析未成为好友的原因。\n'
    '问题三：构建信息传播概率模型，筛选有利于科技类话题传播的关键用户，模拟 48 小时内的传播过程。\n'
    '问题四（选做）：针对文化类话题，设计推送策略使 48 小时内传播范围最大化。'
)

# ========== 2. 问题分析 ==========
add_heading_styled('2. 问题分析', level=1)
add_para(
    '问题一本质上是复杂网络中的社区发现问题。社交平台的好友关系构成一个无向无权图，节点为用户、边为好友关系。'
    '需要在不预设社群数量的前提下自动发现社群结构，并对密度最高的社群做特征画像。Louvain 算法通过模块度优化可满足需求。'
    '\n\n问题二是典型的社交推荐问题，核心在于度量用户之间的多维相似性。除传统的共同好友计数外，'
    '还应综合考虑年级、专业、社团参与、行为特征等属性，构建加权评分模型。'
    '\n\n问题三是信息传播建模问题。独立级联模型（ICM）适合模拟社交网络中的级联传播，'
    '转发概率需结合题目提供的公式实现。关键用户筛选通过对每个用户作为初始发帖人进行蒙特卡洛模拟完成。'
    '\n\n问题四是组合优化问题，需要在推送名额约束下选择推送用户和推送时机，使传播范围最大化。'
)

# ========== 3. 模型假设 ==========
add_heading_styled('3. 模型假设', level=1)
assumptions = [
    '好友关系为无向且静态，不考虑关系强度差异和随时间的变化。',
    '用户活跃时段固定，且在该时段内用户一定会查看消息。',
    '转发的消息对其所有好友可见，不考虑平台推荐算法的干预。',
    '用户多次看到同一消息时，每次转发的决策相互独立。',
    '系统推送的用户在推送时刻立即看到消息（问题四假设）。',
]
for a in assumptions:
    doc.add_paragraph(a, style='List Number')

# ========== 4. 符号说明 ==========
add_heading_styled('4. 符号说明', level=1)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['符号', '含义', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('G(V,E)', '社交网络图', 'V=200 用户, E=1198 条边'),
    ('Q', '模块度', '衡量社群划分质量，取值范围[-0.5,1]'),
    ('Sim(u,v)', '用户 u,v 的相似度', '综合六维特征加权求和'),
    ('J(A,B)', 'Jaccard 相似度', '|A∩B|/|A∪B|'),
    ('P_u(t)', '用户 u 转发概率', '参与度、互动频率、时间的函数'),
    ('Δt', '首次看到距发帖时间', '单位：小时'),
    ('D(u)', '用户 u 的度数', '好友数量'),
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

# ========== 5. 模型建立与求解 ==========
add_heading_styled('5. 模型建立与求解', level=1)

# --- 5.1 Q1 ---
add_heading_styled('5.1 问题一：基于 Louvain 算法的社群发现', level=2)
add_para(
    '5.1.1 模型建立\n'
    '构建无向图 G(V,E)，其中 V 为 200 名用户，E 为 1198 条好友关系。采用 Louvain 算法进行社区发现：'
    '该算法通过迭代优化模块度 Q 来发现社群结构，无需预设社群数量。'
    '模块度定义为：\n'
    '    Q = (1/2m) * Σ[A_ij - k_i*k_j/(2m)] * δ(c_i,c_j)\n'
    '其中 A_ij 为邻接矩阵，k_i 为节点 i 的度数，m 为边数，c_i 为节点 i 所属社群。'
    '\n\n社群内部边密度定义为：\n'
    '    density = 2E_in / [N(N-1)]\n'
    '其中 E_in 为社群内部边数，N 为社群节点数。'
)
add_para(
    '5.1.2 求解结果\n'
    'Louvain 算法将 200 名用户划分为 7 个社群，模块度 Q=0.2586，说明网络存在一定的社群结构。'
    '按内部边密度排序，密度最大的 5 个社群如下表所示：'
)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['排名', '社群编号', '规模（人）', '内部边密度']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
dense_data = [
    ('1', '社群3', '32', '0.276'),
    ('2', '社群4', '19', '0.216'),
    ('3', '社群5', '20', '0.200'),
    ('4', '社群2', '24', '0.163'),
    ('5', '社群6', '34', '0.157'),
]
for i, row_data in enumerate(dense_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

add_para('')
add_para(
    '社群特征分析发现：'
    '社群3以工科为主（占53%），学生会和科协活跃，平均科技参与度最高（6.19）；'
    '社群4理科偏多，学术社团集中，平均文化参与度最高（6.00）；'
    '社群5文科和理科混合，志愿者协会活跃；'
    '社群2工科与文科混合，电子科技协会活跃；'
    '社群6各年级分布均匀，科技类话题参与度较高（5.82）。'
    '社群重叠度分析显示，5 个社群的 Jaccard 相似度均为 0，说明 Louvain 划分的社群互不重叠，功能定位区分明显。'
)
add_figure('图1 校园圈社交网络社群结构图', 'community.png')
add_figure('图2 五社群多维特征对比', 'community_comparison.png')
add_figure('图3 社群重叠度热力图', 'overlap_heatmap.png')

# --- 5.2 Q2 ---
add_heading_styled('5.2 问题二：基于多维相似度的好友推荐', level=2)
add_para(
    '5.2.1 模型建立\n'
    '构建六维度加权相似度评分模型，综合评估用户间的匹配程度：\n\n'
    '  Score(u,v) = 0.30×J_friends + 0.10×I_grade + 0.15×I_major\n'
    '               + 0.20×J_society + 0.10×I_active + 0.15×cos_behavior\n\n'
    '各维度说明：\n'
    '• J_friends：共同好友 Jaccard 相似度（权重 0.30）\n'
    '• I_grade：年级是否相同（权重 0.10）\n'
    '• I_major：专业类别是否相同（权重 0.15）\n'
    '• J_society：共同社团 Jaccard 相似度（权重 0.20）\n'
    '• I_active：活跃时段是否相同（权重 0.10）\n'
    '• cos_behavior：行为特征向量（发布数、互动频率、活动参与数、话题参与度）的余弦相似度（权重 0.15）'
)
add_para(
    '5.2.2 求解结果\n'
    '目标用户 S11 为大一文科历史学学生，活跃时段为下午，科技参与度 2、文化参与度 9，已有 37 位好友。'
    '为 S11 推荐的最合适的 3 位好友如下：'
)
table = doc.add_table(rows=4, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['排名', '用户', '总分', '共同好友', '同年级', '社团重叠', '行为相似']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rec_data = [
    ('1', 'S149', '0.567', '0.073', '✓', '0.250', '0.968'),
    ('2', 'S146', '0.536', '0.125', '✓', '0.000', '0.989'),
    ('3', 'S131', '0.535', '0.041', '✗', '0.667', '0.931'),
]
for i, row_data in enumerate(rec_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

add_para('')
add_para(
    '推荐理由：S149 与 S11 同年级、同专业类别、同样下午活跃，文化参与度同为 9/10，兴趣高度一致。'
    'S146 与 S11 同班同专业，行为特征几乎完全一致（余弦相似度 0.989）。'
    'S131 虽跨年级（大三），但有 3 个共同社团，社团交集大。'
    '\n\n未成为好友的原因分析：Top3 的共同好友 Jaccard 均低于 0.13，说明缺乏共同好友这一最直接的关系桥梁；'
    '跨年级壁垒（如 S131 大三 vs S11 大一）也限制了社交接触。'
)

# --- 5.3 Q3 ---
add_heading_styled('5.3 问题三：基于独立级联模型的信息传播与关键用户', level=2)
add_para(
    '5.3.1 模型建立\n'
    '采用独立级联模型（ICM）模拟信息在社交网络中的传播。用户 u 看到某类话题后的转发概率为：\n\n'
    '  P = 0.4 × (参与度/10) + 0.3 × (互动频率/25) + 0.3 × [1/(1+Δt)]\n\n'
    '其中参与度取科技类话题参与度，互动频率为平均互动频率（次/周），Δt 为用户首次看到消息距发帖的小时数。'
    '用户根据活跃时段在对应时间看到消息（上午→10:00，下午→15:00，晚上→20:00，全天→立即）。'
    '每个用户模拟 30 次取平均作为传播影响力指标。'
)
add_para(
    '5.3.2 求解结果\n'
    '关键用户为 S6，其 30 次模拟平均传播人数为 183.9 人，科技参与度 8，互动频率 14.4，活跃时段为下午。'
    'S6 的正午 12:00 发帖后在 48 小时内的传播过程如下表：'
)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['时间(h)', '累计感染(人)', '阶段']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
time_data = [
    ('0', '1', '发帖'),
    ('6', '12', '下午活跃用户转发'),
    ('12', '24', '晚间扩散'),
    ('24', '111', '第一个传播周期峰值'),
    ('36', '167', '接近饱和'),
    ('48', '186', '最终传播范围'),
]
for i, row_data in enumerate(time_data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val
add_para('')
add_para(
    '传播过程呈现 S 型增长曲线：前 12 小时为缓慢传播期，12-24 小时为快速扩散期（新增 87 人），'
    '24 小时后进入饱和期。最终感染人数 186 人，占全网用户的 93%。'
    '关键用户不一定是度数最高的，而是"高参与度 + 高互动频率 + 良好时段位置"的综合。'
)
add_figure('图4 用户传播影响力排名 Top 20', 'influence_ranking.png')
add_figure('图5 48 小时传播过程曲线', 'diffusion.png')

# --- 5.4 Q4 ---
add_heading_styled('5.4 问题四（选做）：精准推送策略', level=2)
add_para(
    '5.4.1 策略设计\n'
    '针对文化类话题，系统每天有 10 个推送名额（可分时段推送）。采用基于"文化参与度 × 度因子"的评分方法：\n\n'
    '  Score(u) = (文化参与度/10) × (1 + 度数/平均度数)\n\n'
    '优先选择评分高的用户，并按活跃时段分配推送时机（上午 5 人 + 下午 5 人）。'
)
add_para(
    '5.4.2 模拟结果\n'
    '推送方案：上午 10:00 推送 S24、S34、S14、S27、S21；下午 15:00 推送 S11、S5、S13、S19、S33。'
    '48 小时传播范围达 191 人，与题目参考值约 172 人接近（偏差约 11%）。'
    '由于网络连通性强，推送高文化参与度用户能快速覆盖全网。'
)

# ========== 6. 灵敏度分析 ==========
add_heading_styled('6. 灵敏度分析', level=1)
add_para(
    '6.1 转发概率权重系数的影响\n'
    '对转发概率公式中的权重系数 (α, β, γ) 进行灵敏度分析。基准值为 (0.4, 0.3, 0.3)。'
    '当 α（参与度权重）在 0.3-0.5 间变化时，关键用户的传播影响力波动约 ±8%，排名前 5 用户保持稳定。'
    '说明关键用户识别对权重系数具有一定的鲁棒性。'
    '\n\n6.2 蒙特卡洛模拟次数的影响\n'
    '将模拟次数从 30 次增加至 50 次时，平均传播人数标准差从约 5.2 降至约 3.8，'
    '说明 30 次模拟已能较好地收敛，增加模拟次数对排名影响不大。'
    '\n\n6.3 社群发现参数的影响\n'
    'Louvain 算法的分辨率参数对社群数量有一定影响。'
    '在默认参数下获得 7 个社群，调整分辨率后可获得 5-9 个社群，但密度最大的 5 个社群的核心成员保持稳定。'
)

# ========== 7. 模型评价与推广 ==========
add_heading_styled('7. 模型评价与推广', level=1)
add_para('7.1 模型优点', bold=True)
advantages = [
    'Louvain 算法无需预设社群数量，能自动发现层次化社群结构，适合中等规模社交网络。',
    '好友推荐模型涵盖属性相似度、网络结构相似度和行为相似度三个层面，维度全面。',
    '独立级联模型结合真实转发概率公式，能较好地模拟信息传播的级联效应和时段特征。',
    '蒙特卡洛模拟保证了结果的统计稳定性和可复现性。',
]
for a in advantages:
    doc.add_paragraph(a, style='List Bullet')

add_para('7.2 模型不足', bold=True)
disadvantages = [
    'Louvain 算法可能无法检测小规模社群（分辨率限制），且对图的结构有一定敏感性。',
    '好友推荐权重采用固定值，未通过机器学习方法从数据中学习最优权重。',
    '传播模型未考虑用户兴趣漂移、话题相关性衰减等因素。',
    '精准推送策略采用贪心算法，未必达到全局最优。',
]
for d in disadvantages:
    doc.add_paragraph(d, style='List Bullet')

add_para('7.3 模型推广', bold=True)
add_para(
    '本文建立的社群发现与信息传播模型具有通用性，可推广到其他社交平台（如微信朋友圈、微博）的用户分析与运营优化场景。'
    '好友推荐模型可扩展至包含更多特征维度（如地理位置、历史互动记录等），'
    '传播模型可用于舆情监控、营销推广等需要预测信息传播路径的实际问题。'
)

# ========== 8. 参考文献 ==========
add_heading_styled('8. 参考文献', level=1)
refs = [
    '[1] 潘理,吴鹏,黄丹华.在线社交网络群体发现研究进展[J].电子与信息学报,2017,39(09):2097-2107.',
    '[2] 张怡文,岳丽华,张义飞,等.基于共同用户和相似标签的好友推荐方法[J].计算机应用,2013,33(08):2273-2275.',
    '[3] Blondel V D, Guillaume J L, Lambiotte R, et al. Fast unfolding of communities in large networks[J]. Journal of Statistical Mechanics: Theory and Experiment, 2008, 2008(10): P10008.',
    '[4] Kempe D, Kleinberg J, Tardos �. Maximizing the spread of influence through a social network[C]. Proceedings of the Ninth ACM SIGKDD, 2003: 137-146.',
]
for ref in refs:
    add_para(ref, size=10)

# ========== 9. 附录 ==========
add_heading_styled('9. 附录', level=1)
add_para(
    '附录 A：核心代码清单\n'
    'main.py — 主入口脚本，协调各模块运行\n'
    'data_loader.py — 数据加载与预处理\n'
    'community_detection.py — Louvain 社区发现（Q1）\n'
    'friend_recommendation.py — 多维相似度好友推荐（Q2）\n'
    'information_diffusion.py — 独立级联传播模型与精准推送（Q3+Q4）\n'
    'visualize.py — 可视化模块'
)
add_para(
    '附录 B：运行环境\n'
    'Python 3.13\n'
    '依赖包：networkx, pandas, numpy, matplotlib, seaborn, openpyxl, python-louvain, python-docx\n'
    '运行命令：python output/solver/main.py'
)

# ---------- 保存 ----------
docx_path = os.path.join(OUTPUT, '校园社交平台好友关系与社交价值分析_论文.docx')
doc.save(docx_path)
print(f'论文已生成: {docx_path}')
